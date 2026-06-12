"""结构感知文档解析层（V2.0 IDP-01）。

V1.5 的 parse_document() 输出纯文本 `str`；V2.0 升级为输出 `list[StructuredBlock]`，
每个 block 携带 block_type / heading_path / page_number / position_index 等元数据。

支持格式：
- .pdf  → PyMuPDF 提取 text blocks，按字体大小/粗细推断标题层级（启发式）
- .docx → python-docx 直接读 paragraph.style.name，匹配 "Heading 1/2/3"
- .md   → markdown-it-py 的 token 类型直接对应
- .txt  → 单 block_type=paragraph，heading_path=[]

设计要点：
- 保留 V1.5 的 `parse_document()` 作为兼容入口（返回 str），供 V1.5 /api/v1/ 使用
- 新增 `parse_document_structured()` 返回 `list[StructuredBlock]`，供 V2.0 使用
- 两者共享底层解析器，只差最后一步拼装
"""

from __future__ import annotations

import logging
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Callable, Literal

logger = logging.getLogger(__name__)


# ──────────────── 公共类型 ────────────────


# V1.5 兼容：支持的扩展名
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".pdf", ".docx", ".md", ".txt"})

# 各扩展名期望的 MIME 集合
EXPECTED_MIMES: dict[str, frozenset[str]] = {
    ".pdf": frozenset({"application/pdf"}),
    ".docx": frozenset(
        {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
    ),
    ".md": frozenset({"text/markdown", "text/x-markdown", "text/plain"}),
    ".txt": frozenset({"text/plain"}),
}


class ParseError(Exception):
    """解析失败的统一异常。"""


# ──────────────── V2.0 结构感知类型 ────────────────


@dataclass(frozen=True)
class StructuredBlock:
    """V2.0 结构感知解析的输出单元（IDP-01）。

    每个块代表文档中一个语义完整的片段，携带结构元数据。
    后续 StructuredSplitter 会基于这些元数据做智能切片。
    """

    block_id: str  # uuid，全局唯一
    block_type: Literal["paragraph", "heading", "table", "code", "list"]
    heading_path: list[str]  # 该 block 所属的标题层级路径，如 ["第1章", "1.1 节"]
    content: str  # 块内容（表格类型为 markdown 格式）
    page_number: int | None  # 页码（PDF 有效）
    position_index: int  # 文档内序号，从 0 开始


def _make_block_id() -> str:
    """生成唯一 block_id。"""
    return uuid.uuid4().hex


# ──────────────── V2.0 结构感知解析器 ────────────────


def _parse_pdf_structured(path: Path) -> list[StructuredBlock]:
    """PDF 结构感知解析：按字体大小/粗细推断标题层级。

    启发式规则：
    - 统计全文所有 text block 的字号分布
    - 大字号（> 中位数 + 阈值）且粗体 → 标题
    - 其余 → 段落
    - 表格检测：文本含多行 tab/| 分隔符 → 标记为 table
    """
    try:
        import fitz  # pymupdf
    except ImportError as e:
        raise ParseError("解析 PDF 需要 pymupdf 库") from e

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        raise ParseError(f"PDF 打开失败：{e}") from e

    blocks: list[StructuredBlock] = []
    position = 0

    try:
        # 第一遍：收集所有 text block 的字号信息，用于标题推断
        all_font_sizes: list[float] = []
        page_blocks_raw: list[tuple[int, str, float, bool]] = []  # (page, text, fontsize, bold)

        for page_idx, page in enumerate(doc):
            text_dict = page.get_text("dict")
            for block_info in text_dict.get("blocks", []):
                if block_info.get("type") != 0:  # 0=text block, 1=image block
                    continue
                block_text_parts: list[str] = []
                max_font_size = 0.0
                is_bold = False
                for line in block_info.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        block_text_parts.append(text)
                        font_size = span.get("size", 12.0)
                        max_font_size = max(max_font_size, font_size)
                        # 粗体检测：font name 含 Bold/粗体
                        font_name = span.get("font", "").lower()
                        if "bold" in font_name or "黑体" in font_name or "heavy" in font_name:
                            is_bold = True

                block_text = " ".join(block_text_parts).strip()
                if not block_text:
                    continue
                all_font_sizes.append(max_font_size)
                page_blocks_raw.append((page_idx + 1, block_text, max_font_size, is_bold))

        # 计算字号阈值用于标题推断
        if all_font_sizes:
            sorted_sizes = sorted(all_font_sizes)
            median_size = sorted_sizes[len(sorted_sizes) // 2]
            # 标题判定：字号 > 中位数 * 1.2 且粗体；或者字号 > 中位数 * 1.5（即使不粗体）
            heading_threshold = median_size * 1.2
            heading_strong_threshold = median_size * 1.5
        else:
            heading_threshold = 14.0
            heading_strong_threshold = 18.0

        # 标题层级推断：按字号大小映射到 h1/h2/h3
        # 找出所有可能的标题字号（去重排序）
        heading_sizes = sorted(
            set(
                s for s in all_font_sizes
                if s > heading_threshold or (s > heading_strong_threshold)
            ),
            reverse=True,
        )

        def _infer_heading_level(size: float, bold: bool) -> int | None:
            """根据字号和粗体推断标题层级（1/2/3），非标题返回 None。"""
            if not bold and size <= heading_strong_threshold:
                return None
            if bold and size <= heading_threshold:
                return None
            # 映射字号到层级
            for level, size_threshold in enumerate(heading_sizes, 1):
                if size >= size_threshold * 0.95:  # 容差 5%
                    # 层级最大 3
                    return min(level, 3)
            return None

        # 第二遍：构建 StructuredBlock
        current_heading_path: list[str] = []
        current_heading_levels: dict[int, str] = {}  # level → title

        for page_num, text, font_size, bold in page_blocks_raw:
            heading_level = _infer_heading_level(font_size, bold)

            if heading_level is not None:
                # 这是一个标题 block
                # 更新标题路径：清除当前层级及以下的标题
                for lvl in list(current_heading_levels.keys()):
                    if lvl >= heading_level:
                        del current_heading_levels[lvl]
                current_heading_levels[heading_level] = text
                # 重建 heading_path
                current_heading_path = [
                    current_heading_levels[lvl]
                    for lvl in sorted(current_heading_levels.keys())
                ]

                block = StructuredBlock(
                    block_id=_make_block_id(),
                    block_type="heading",
                    heading_path=list(current_heading_path[:-1]),  # 不含自身
                    content=text,
                    page_number=page_num,
                    position_index=position,
                )
            else:
                # 判断是否为表格（启发式：含多个 | 分隔符或多行 tab 分隔）
                is_table = _detect_table(text)
                block = StructuredBlock(
                    block_id=_make_block_id(),
                    block_type="table" if is_table else "paragraph",
                    heading_path=list(current_heading_path),
                    content=text,
                    page_number=page_num,
                    position_index=position,
                )

            blocks.append(block)
            position += 1

    finally:
        doc.close()

    if not blocks:
        logger.warning("PDF 解析: 未提取到任何结构化内容 path=%s", path)

    return blocks


def _detect_table(text: str) -> bool:
    """启发式判断文本是否为表格内容。

    判定规则（任一满足）：
    - 连续 3 行以上含 | 分隔符（markdown 表格风格）
    - 含多行 tab 分隔且列数一致
    """
    lines = text.strip().split("\n")
    pipe_lines = sum(1 for line in lines if "|" in line and line.strip().count("|") >= 2)
    if pipe_lines >= 3:
        return True
    tab_lines = sum(1 for line in lines if "\t" in line and line.count("\t") >= 2)
    if tab_lines >= 2:
        return True
    return False


def _parse_docx_structured(path: Path) -> list[StructuredBlock]:
    """DOCX 结构感知解析：直接读取 paragraph.style.name 匹配标题层级。"""
    try:
        from docx import Document
    except ImportError as e:
        raise ParseError("解析 docx 需要 python-docx 库") from e

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ParseError(f"docx 打开失败：{e}") from e

    blocks: list[StructuredBlock] = []
    position = 0
    current_heading_path: list[str] = []
    current_heading_levels: dict[int, str] = {}

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = (para.style.name or "").lower() if para.style else ""
        heading_level = _parse_docx_heading_level(style_name)

        if heading_level is not None:
            # 标题段落
            for lvl in list(current_heading_levels.keys()):
                if lvl >= heading_level:
                    del current_heading_levels[lvl]
            current_heading_levels[heading_level] = text
            current_heading_path = [
                current_heading_levels[lvl]
                for lvl in sorted(current_heading_levels.keys())
            ]

            blocks.append(
                StructuredBlock(
                    block_id=_make_block_id(),
                    block_type="heading",
                    heading_path=list(current_heading_path[:-1]),
                    content=text,
                    page_number=None,
                    position_index=position,
                )
            )
        else:
            # 普通段落
            blocks.append(
                StructuredBlock(
                    block_id=_make_block_id(),
                    block_type="paragraph",
                    heading_path=list(current_heading_path),
                    content=text,
                    page_number=None,
                    position_index=position,
                )
            )
        position += 1

    # 表格处理：每个表格转为 markdown 格式
    for table in doc.tables:
        rows_text: list[str] = []
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            rows_text.append("| " + " | ".join(cells) + " |")
            # 第一行后加分隔线（markdown 表格格式）
            if row_idx == 0:
                rows_text.append("| " + " | ".join("---" for _ in cells) + " |")

        table_md = "\n".join(rows_text)
        blocks.append(
            StructuredBlock(
                block_id=_make_block_id(),
                block_type="table",
                heading_path=list(current_heading_path),
                content=table_md,
                page_number=None,
                position_index=position,
            )
        )
        position += 1

    return blocks


def _parse_docx_heading_level(style_name: str) -> int | None:
    """从 docx paragraph style name 推断标题层级。

    匹配规则：
    - "heading 1" / "heading1" / "标题 1" → 1
    - "heading 2" / "标题 2" → 2
    - "heading 3" / "标题 3" → 3
    - 其他 → None
    """
    import re

    # 英文风格：Heading 1 / Heading1
    m = re.search(r"heading\s*(\d+)", style_name)
    if m:
        return min(int(m.group(1)), 3)
    # 中文风格：标题 1 / 标题1
    m = re.search(r"标题\s*(\d+)", style_name)
    if m:
        return min(int(m.group(1)), 3)
    return None


def _parse_md_structured(path: Path) -> list[StructuredBlock]:
    """Markdown 结构感知解析：markdown-it-py token 类型直接对应。"""
    try:
        from markdown_it import MarkdownIt
    except ImportError as e:
        raise ParseError("解析 markdown 需要 markdown-it-py 库") from e

    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as e:
        raise ParseError(f"markdown 文件非 UTF-8 编码：{e}") from e

    md = MarkdownIt().enable("table")
    tokens = md.parse(raw)

    blocks: list[StructuredBlock] = []
    position = 0
    current_heading_path: list[str] = []
    current_heading_levels: dict[int, str] = {}

    i = 0
    while i < len(tokens):
        tok = tokens[i]

        # ── 标题 ──
        if tok.type == "heading_open":
            # heading_open 后紧跟 inline，再 heading_close
            level = _parse_md_heading_level(tok.tag)
            # 取下一个 inline token 的内容作为标题文本
            heading_text = ""
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                heading_text = tokens[i + 1].content.strip()

            if level is not None and heading_text:
                for lvl in list(current_heading_levels.keys()):
                    if lvl >= level:
                        del current_heading_levels[lvl]
                current_heading_levels[level] = heading_text
                current_heading_path = [
                    current_heading_levels[lvl]
                    for lvl in sorted(current_heading_levels.keys())
                ]
                blocks.append(
                    StructuredBlock(
                        block_id=_make_block_id(),
                        block_type="heading",
                        heading_path=list(current_heading_path[:-1]),
                        content=heading_text,
                        page_number=None,
                        position_index=position,
                    )
                )
                position += 1
            i += 3  # 跳过 heading_open + inline + heading_close
            continue

        # ── 代码块（fence / code_block）──
        if tok.type in ("fence", "code_block"):
            if tok.content.strip():
                blocks.append(
                    StructuredBlock(
                        block_id=_make_block_id(),
                        block_type="code",
                        heading_path=list(current_heading_path),
                        content=tok.content,
                        page_number=None,
                        position_index=position,
                    )
                )
                position += 1
            i += 1
            continue

        # ── 表格 ──
        if tok.type == "table_open":
            table_md = _extract_md_table(tokens, i)
            if table_md:
                blocks.append(
                    StructuredBlock(
                        block_id=_make_block_id(),
                        block_type="table",
                        heading_path=list(current_heading_path),
                        content=table_md,
                        page_number=None,
                        position_index=position,
                    )
                )
                position += 1
            # 跳过整个 table_open ... table_close 序列
            while i < len(tokens) and tokens[i].type != "table_close":
                i += 1
            i += 1  # 跳过 table_close
            continue

        # ── 段落 ──
        if tok.type == "paragraph_open":
            text = ""
            if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                text = tokens[i + 1].content.strip()
            if text:
                blocks.append(
                    StructuredBlock(
                        block_id=_make_block_id(),
                        block_type="paragraph",
                        heading_path=list(current_heading_path),
                        content=text,
                        page_number=None,
                        position_index=position,
                    )
                )
                position += 1
            i += 3  # paragraph_open + inline + paragraph_close
            continue

        # ── 列表项 ──
        if tok.type == "bullet_list_open" or tok.type == "ordered_list_open":
            list_items = _extract_md_list(tokens, i)
            if list_items:
                blocks.append(
                    StructuredBlock(
                        block_id=_make_block_id(),
                        block_type="list",
                        heading_path=list(current_heading_path),
                        content=list_items,
                        page_number=None,
                        position_index=position,
                    )
                )
                position += 1
            # 跳过整个列表
            close_type = tok.type.replace("_open", "_close")
            while i < len(tokens) and tokens[i].type != close_type:
                i += 1
            i += 1
            continue

        i += 1

    return blocks


def _parse_md_heading_level(tag: str) -> int | None:
    """从 heading_open token 的 tag 属性推断层级。h1→1, h2→2, h3→3。"""
    import re

    m = re.match(r"h(\d+)", tag)
    if m:
        return min(int(m.group(1)), 3)
    return None


def _extract_md_table(tokens: list, start: int) -> str:
    """从 markdown-it token 流中提取表格的 markdown 文本。"""
    # 简化处理：直接从 token 的内容重建 markdown 表格
    rows: list[str] = []
    i = start + 1  # 跳过 table_open
    while i < len(tokens) and tokens[i].type != "table_close":
        tok = tokens[i]
        if tok.type == "tr_open":
            cells: list[str] = []
            i += 1
            while i < len(tokens) and tokens[i].type != "tr_close":
                if tokens[i].type == "td_open" or tokens[i].type == "th_open":
                    # 下一个应该是 inline
                    if i + 1 < len(tokens) and tokens[i + 1].type == "inline":
                        cells.append(tokens[i + 1].content.strip())
                i += 1
            if cells:
                row_md = "| " + " | ".join(cells) + " |"
                if not rows:
                    # 第一行后加分隔线
                    rows.append(row_md)
                    rows.append("| " + " | ".join("---" for _ in cells) + " |")
                else:
                    rows.append(row_md)
        i += 1
    return "\n".join(rows)


def _extract_md_list(tokens: list, start: int) -> str:
    """从 markdown-it token 流中提取列表项文本。"""
    items: list[str] = []
    i = start + 1
    close_type = tokens[start].type.replace("_open", "_close")
    while i < len(tokens) and tokens[i].type != close_type:
        tok = tokens[i]
        if tok.type == "list_item_open":
            # 收集该 item 的 inline 内容
            i += 1
            item_parts: list[str] = []
            while i < len(tokens) and tokens[i].type != "list_item_close":
                if tokens[i].type == "inline":
                    item_parts.append(tokens[i].content.strip())
                elif tokens[i].type == "paragraph_close":
                    item_parts.append("\n")
                i += 1
            if item_parts:
                items.append("- " + " ".join(item_parts).strip())
        i += 1
    return "\n".join(items)


def _parse_txt_structured(path: Path) -> list[StructuredBlock]:
    """纯文本解析：每段为单个 paragraph block，heading_path=[]。"""
    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as e:
        raise ParseError(f"txt 文件非 UTF-8 编码：{e}") from e

    # 按空行分段
    paragraphs = [p.strip() for p in raw.split("\n\n") if p.strip()]
    if not paragraphs:
        # 退化：按行分段
        paragraphs = [line.strip() for line in raw.split("\n") if line.strip()]

    blocks: list[StructuredBlock] = []
    for i, para in enumerate(paragraphs):
        blocks.append(
            StructuredBlock(
                block_id=_make_block_id(),
                block_type="paragraph",
                heading_path=[],
                content=para,
                page_number=None,
                position_index=i,
            )
        )

    return blocks


# ──────────────── 分发表 ────────────────


# V2.0 结构感知解析器
_STRUCTURED_PARSERS: dict[str, Callable[[Path], list[StructuredBlock]]] = {
    ".pdf": _parse_pdf_structured,
    ".docx": _parse_docx_structured,
    ".md": _parse_md_structured,
    ".txt": _parse_txt_structured,
}

# V1.5 兼容解析器（返回 str；保留原始 V1.5 实现，不包装结构化解析器）
def _parse_pdf_text(path: Path) -> str:
    """V1.5 兼容：PDF → 纯文本。原始实现保留，空页 warning 逻辑不变。"""
    try:
        import fitz  # pymupdf
    except ImportError as e:
        raise ParseError("解析 PDF 需要 pymupdf 库") from e

    parts: list[str] = []
    try:
        doc = fitz.open(str(path))
    except Exception as e:
        raise ParseError(f"PDF 打开失败：{e}") from e

    try:
        empty_pages: list[int] = []
        for page_idx, page in enumerate(doc):
            text = page.get_text() or ""
            if not text.strip():
                empty_pages.append(page_idx + 1)
                continue
            parts.append(text)
        if empty_pages:
            logger.warning(
                "PDF 解析: %d 页无文本(可能是扫描页)已跳过 - 页码=%s",
                len(empty_pages),
                empty_pages,
            )
    finally:
        doc.close()

    return "\n\n".join(parts)


def _parse_docx_text(path: Path) -> str:
    """V1.5 兼容：DOCX → 纯文本。"""
    try:
        from docx import Document
    except ImportError as e:
        raise ParseError("解析 docx 需要 python-docx 库") from e

    try:
        doc = Document(str(path))
    except Exception as e:
        raise ParseError(f"docx 打开失败：{e}") from e

    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n\n".join(parts)


def _parse_md_text(path: Path) -> str:
    """V1.5 兼容：Markdown 剥语法符号后取纯文本。"""
    try:
        from markdown_it import MarkdownIt
    except ImportError as e:
        raise ParseError("解析 markdown 需要 markdown-it-py 库") from e

    try:
        raw = path.read_text(encoding="utf-8")
    except UnicodeDecodeError as e:
        raise ParseError(f"markdown 文件非 UTF-8 编码：{e}") from e

    md = MarkdownIt()
    tokens = md.parse(raw)

    parts: list[str] = []

    def _walk(toks) -> None:
        for tok in toks:
            if tok.type == "inline" and tok.children:
                _walk(tok.children)
                continue
            if tok.type in ("text", "code_inline", "code_block", "fence"):
                if tok.content:
                    parts.append(tok.content)
            if tok.type in (
                "paragraph_close",
                "heading_close",
                "list_item_close",
                "fence",
                "code_block",
            ):
                parts.append("\n")

    _walk(tokens)

    joined = "".join(parts)
    while "\n\n\n" in joined:
        joined = joined.replace("\n\n\n", "\n\n")
    return joined.strip()


def _parse_txt_text(path: Path) -> str:
    """V1.5 兼容：TXT → 纯文本。"""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError as e:
        raise ParseError(f"txt 文件非 UTF-8 编码：{e}") from e


_PARSERS = {
    ".pdf": _parse_pdf_text,
    ".docx": _parse_docx_text,
    ".md": _parse_md_text,
    ".txt": _parse_txt_text,
}


# ──────────────── 对外入口 ────────────────


def is_supported_filename(filename: str) -> bool:
    """按扩展名判断是否支持。"""
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def check_mime_compatibility(filename: str, declared_mime: str | None) -> bool:
    """二次校验：检查声明的 MIME 是否与扩展名匹配。"""
    if not declared_mime:
        return True
    ext = Path(filename).suffix.lower()
    expected = EXPECTED_MIMES.get(ext, frozenset())
    return declared_mime.lower() in expected


def parse_document(path: str | Path, *, filename: str | None = None) -> str:
    """V1.5 兼容入口：根据扩展名分发到具体解析器，返回纯文本。

    Args:
        path: 磁盘路径
        filename: 原始文件名（用来取扩展名）；默认从 path 取

    Returns:
        解析后的纯文本

    Raises:
        ParseError: 文件不存在 / 格式不支持 / 解析失败
    """
    p = Path(path)
    if not p.exists():
        raise ParseError(f"文件不存在: {p}")

    ext = Path(filename or p.name).suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise ParseError(
            f"不支持的文件扩展名: {ext}（当前支持 {sorted(SUPPORTED_EXTENSIONS)}）"
        )

    logger.info("解析文件(V1.5兼容): name=%s ext=%s", filename or p.name, ext)
    text = parser(p)
    logger.info("解析完成: name=%s 文本长度=%d", filename or p.name, len(text))
    return text


def parse_document_structured(
    path: str | Path, *, filename: str | None = None
) -> list[StructuredBlock]:
    """V2.0 结构感知解析入口：返回 StructuredBlock 列表。

    Args:
        path: 磁盘路径
        filename: 原始文件名（用来取扩展名）；默认从 path 取

    Returns:
        StructuredBlock 列表，按文档出现顺序排列

    Raises:
        ParseError: 文件不存在 / 格式不支持 / 解析失败
    """
    p = Path(path)
    if not p.exists():
        raise ParseError(f"文件不存在: {p}")

    ext = Path(filename or p.name).suffix.lower()
    parser = _STRUCTURED_PARSERS.get(ext)
    if parser is None:
        raise ParseError(
            f"不支持的文件扩展名: {ext}（当前支持 {sorted(SUPPORTED_EXTENSIONS)}）"
        )

    logger.info("结构感知解析: name=%s ext=%s", filename or p.name, ext)
    blocks = parser(p)
    logger.info(
        "结构感知解析完成: name=%s blocks=%d types=%s",
        filename or p.name,
        len(blocks),
        {b.block_type for b in blocks},
    )
    return blocks


__all__ = [
    # V1.5 兼容
    "SUPPORTED_EXTENSIONS",
    "EXPECTED_MIMES",
    "ParseError",
    "is_supported_filename",
    "check_mime_compatibility",
    "parse_document",
    # V2.0 结构感知
    "StructuredBlock",
    "parse_document_structured",
]
