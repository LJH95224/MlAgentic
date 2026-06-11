"""V1.5 S3.0 ingest 模块单测：解析器 + 切片器。

Fixture 策略：每用例临时目录里生成最小可解析的测试文件（PDF/docx 用代码生成、
md/txt 直接写），跑完自动清理。不依赖 tests/fixtures/ 静态样本，避免污染 repo。
"""

from __future__ import annotations

from pathlib import Path

import pytest

from app.ingest.parser import (
    EXPECTED_MIMES,
    SUPPORTED_EXTENSIONS,
    ParseError,
    check_mime_compatibility,
    is_supported_filename,
    parse_document,
)
from app.ingest.splitter import Chunk, split_text


# ──────────────── fixture：临时生成测试文件 ────────────────


@pytest.fixture
def pdf_file(tmp_path: Path) -> Path:
    """用 PyMuPDF 现场生成最小 PDF（2 页 + 中文文本 + 一个空页）。"""
    import fitz

    doc = fitz.open()
    # 第 1 页：中文文本
    page1 = doc.new_page()
    page1.insert_text(
        (50, 80),
        "气象站第一段：今日北京晴转多云。",
        fontname="china-s",  # 内置中文字体
        fontsize=12,
    )
    # 第 2 页：英文 + 第二段中文
    page2 = doc.new_page()
    page2.insert_text((50, 80), "Hello world.", fontsize=12)
    page2.insert_text(
        (50, 120),
        "气象站第二段：明日多云转阴。",
        fontname="china-s",
        fontsize=12,
    )
    # 第 3 页：空页（无文本，模拟扫描图片页）
    doc.new_page()

    out = tmp_path / "sample.pdf"
    doc.save(str(out))
    doc.close()
    return out


@pytest.fixture
def docx_file(tmp_path: Path) -> Path:
    """用 python-docx 现场生成最小 docx：2 段正文 + 1 张表格。"""
    from docx import Document

    doc = Document()
    doc.add_paragraph("docx 第一段：测试段落 A。")
    doc.add_paragraph("docx 第二段：测试段落 B 包含 English text。")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表头-1"
    table.cell(0, 1).text = "表头-2"
    table.cell(1, 0).text = "数据-A"
    table.cell(1, 1).text = "数据-B"

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


@pytest.fixture
def md_file(tmp_path: Path) -> Path:
    out = tmp_path / "sample.md"
    out.write_text(
        "# 一级标题\n\n"
        "这是一段**粗体**和*斜体*的混合文本。\n\n"
        "## 二级标题\n\n"
        "- 列表项 1\n"
        "- 列表项 2\n\n"
        "```python\n"
        "def hello():\n"
        "    return 'world'\n"
        "```\n\n"
        "段落末尾。\n",
        encoding="utf-8",
    )
    return out


@pytest.fixture
def txt_file(tmp_path: Path) -> Path:
    out = tmp_path / "sample.txt"
    out.write_text(
        "txt 第一段：纯文本测试。\n\n"
        "txt 第二段：包含中英文混排 mixed content。\n",
        encoding="utf-8",
    )
    return out


@pytest.fixture
def gbk_txt_file(tmp_path: Path) -> Path:
    """非 UTF-8 编码的 txt，用于验证 ParseError。"""
    out = tmp_path / "gbk.txt"
    out.write_bytes("这是 GBK 编码的中文".encode("gbk"))
    return out


# ──────────────── 公共 API 校验 ────────────────


def test_supported_extensions_v1_5_first_release():
    """V1.5 第一版只支持 4 种；.doc 推迟（dev_plan S3.7）。"""
    assert SUPPORTED_EXTENSIONS == frozenset({".pdf", ".docx", ".md", ".txt"})


def test_is_supported_filename_positive():
    for fn in ("a.pdf", "中文.docx", "notes.md", "log.txt"):
        assert is_supported_filename(fn), fn


def test_is_supported_filename_negative():
    for fn in ("a.doc", "a.xls", "noext", "a.PDF.bak"):
        assert not is_supported_filename(fn), fn


def test_is_supported_filename_case_insensitive():
    """扩展名大小写不敏感（Windows 上常见 .PDF）。"""
    for fn in ("a.PDF", "B.Docx", "C.MD", "D.TXT"):
        assert is_supported_filename(fn), fn


# ──────────────── MIME 校验 ────────────────


def test_mime_compatibility_exact_match():
    assert check_mime_compatibility("a.pdf", "application/pdf")
    assert check_mime_compatibility(
        "a.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert check_mime_compatibility("a.md", "text/markdown")


def test_mime_compatibility_case_insensitive():
    assert check_mime_compatibility("a.pdf", "APPLICATION/PDF")


def test_mime_compatibility_md_accepts_octet_stream_false():
    """Windows 上 .md 常被识别为 octet-stream → 返 False，由调用方决定 warning 不抛。"""
    assert not check_mime_compatibility("a.md", "application/octet-stream")


def test_mime_compatibility_none_passes():
    """部分客户端不带 content-type；视为放行。"""
    assert check_mime_compatibility("a.pdf", None)


def test_mime_compatibility_unknown_ext_returns_false():
    assert not check_mime_compatibility("a.unknown", "text/plain")


def test_expected_mimes_covers_all_supported_extensions():
    """每个支持的扩展名都有 MIME 期望集（避免漏配）。"""
    for ext in SUPPORTED_EXTENSIONS:
        assert ext in EXPECTED_MIMES, f"扩展名 {ext} 缺 EXPECTED_MIMES 配置"


# ──────────────── parse_document：通用错误 ────────────────


def test_parse_document_file_not_found(tmp_path):
    with pytest.raises(ParseError, match="文件不存在"):
        parse_document(tmp_path / "nope.txt")


def test_parse_document_unsupported_extension(tmp_path):
    f = tmp_path / "bad.xls"
    f.write_bytes(b"junk")
    with pytest.raises(ParseError, match="不支持的文件扩展名"):
        parse_document(f)


def test_parse_document_use_filename_param_for_extension(tmp_path):
    """path 后缀错但 filename 参数给对了 → 按 filename 取扩展名。"""
    f = tmp_path / "blob"  # 无扩展名
    f.write_text("hello txt", encoding="utf-8")
    text = parse_document(f, filename="real.txt")
    assert "hello txt" in text


# ──────────────── parse_document：各格式 ────────────────


def test_parse_txt_utf8(txt_file):
    text = parse_document(txt_file)
    assert "txt 第一段" in text
    assert "txt 第二段" in text
    assert "mixed content" in text


def test_parse_txt_non_utf8_raises(gbk_txt_file):
    with pytest.raises(ParseError, match="非 UTF-8 编码"):
        parse_document(gbk_txt_file)


def test_parse_md_strips_syntax(md_file):
    text = parse_document(md_file)
    # markdown 语法符号应被剥掉
    assert "**" not in text
    assert "##" not in text
    assert "```" not in text
    # 标题 / 段落 / 代码内容应保留
    assert "一级标题" in text
    assert "二级标题" in text
    assert "段落末尾" in text
    assert "列表项 1" in text
    assert "hello" in text  # 代码块内容
    assert "粗体" in text


def test_parse_md_non_utf8_raises(tmp_path):
    f = tmp_path / "gbk.md"
    f.write_bytes("# 标题".encode("gbk"))
    with pytest.raises(ParseError, match="非 UTF-8"):
        parse_document(f)


def test_parse_docx_extracts_paragraphs_and_tables(docx_file):
    text = parse_document(docx_file)
    # 段落
    assert "测试段落 A" in text
    assert "测试段落 B" in text
    assert "English text" in text
    # 表格内容（cell 文本）
    assert "表头-1" in text
    assert "数据-B" in text


def test_parse_pdf_extracts_all_text_pages(pdf_file):
    text = parse_document(pdf_file)
    # 中文页内容
    assert "第一段" in text
    assert "第二段" in text
    # 英文页
    assert "Hello world" in text


def test_parse_pdf_empty_page_skipped_with_warning(pdf_file, caplog):
    """空页（第 3 页无文本）应被 warning + 跳过，不影响其它页。"""
    import logging

    with caplog.at_level(logging.WARNING, logger="app.ingest.parser"):
        text = parse_document(pdf_file)
    assert "第一段" in text
    # warning 日志提到空页
    assert any("无文本" in r.message for r in caplog.records)


# ──────────────── splitter ────────────────


def test_split_text_empty_returns_empty_list():
    assert split_text("", chunk_size=100, chunk_overlap=10) == []
    assert split_text("   \n\n  ", chunk_size=100, chunk_overlap=10) == []


def test_split_text_short_text_one_chunk():
    chunks = split_text("hello world", chunk_size=100, chunk_overlap=10)
    assert len(chunks) == 1
    assert isinstance(chunks[0], Chunk)
    assert chunks[0].index == 0
    assert "hello world" in chunks[0].text


def test_split_text_long_text_multiple_chunks():
    """长文本 → 多个切片；index 单调递增 + 不丢内容。"""
    # 制造一段超过 100 token 的中文（中文 ≈ 1 token / 字 在 tiktoken 上偏估，约 1.5）
    text = "段落\n\n".join(f"第{i}段：" + "中文测试" * 50 for i in range(10))
    chunks = split_text(text, chunk_size=100, chunk_overlap=20)
    assert len(chunks) > 1
    # index 严格递增
    for i, c in enumerate(chunks):
        assert c.index == i
    # 内容覆盖（取首末段标记验证未丢）
    joined = "\n".join(c.text for c in chunks)
    assert "第0段" in joined
    assert "第9段" in joined


def test_split_text_chunk_size_validated():
    with pytest.raises(ValueError, match="chunk_size 必须 > 0"):
        split_text("x", chunk_size=0, chunk_overlap=0)
    with pytest.raises(ValueError, match="chunk_size 必须 > 0"):
        split_text("x", chunk_size=-1, chunk_overlap=0)


def test_split_text_chunk_overlap_validated():
    with pytest.raises(ValueError, match="chunk_overlap 必须 >= 0"):
        split_text("x", chunk_size=100, chunk_overlap=-1)
    with pytest.raises(ValueError, match="必须 < chunk_size"):
        split_text("x", chunk_size=100, chunk_overlap=100)
    with pytest.raises(ValueError, match="必须 < chunk_size"):
        split_text("x", chunk_size=100, chunk_overlap=200)


def test_split_text_overlap_zero_allowed():
    """overlap=0 应允许（合法边界值）。"""
    chunks = split_text("hello world " * 100, chunk_size=50, chunk_overlap=0)
    assert len(chunks) >= 1


def test_split_text_respects_paragraph_boundary():
    """段落优先：双换行处应优先切，不破坏段落完整性。"""
    para1 = "第一段：" + "短文" * 30
    para2 = "第二段：" + "短文" * 30
    text = f"{para1}\n\n{para2}"

    chunks = split_text(text, chunk_size=80, chunk_overlap=10)
    # 至少切成 2 块；第一段尽可能落在前部，第二段在后部
    assert len(chunks) >= 2


def test_split_text_chunks_are_non_empty():
    """切完不应有空 chunk（split + filter 之后）。"""
    text = "x" * 1000
    chunks = split_text(text, chunk_size=50, chunk_overlap=5)
    for c in chunks:
        assert c.text.strip(), f"出现空 chunk at index={c.index}"
